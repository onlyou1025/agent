"""
文档解析模块
支持 PDF 和 Word 文档的文本提取
"""
import logging
from pathlib import Path
from typing import Optional
import PyPDF2
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class DocumentParser:
    """文档解析器"""
    
    @staticmethod
    def parse_pdf(file_path: str) -> Optional[str]:
        """
        解析 PDF 文件，提取文本内容
        
        Args:
            file_path: PDF 文件路径
            
        Returns:
            提取的文本内容，失败时返回 None
        """
        try:
            text_content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                logger.info(f"开始解析 PDF: {file_path}, 共 {total_pages} 页")
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        text = page.extract_text()
                        if text:
                            text_content.append(f"--- 第 {page_num} 页 ---\n{text}")
                    except Exception as e:
                        logger.warning(f"解析第 {page_num} 页时出错: {str(e)}")
                        continue
            
            full_text = "\n\n".join(text_content)
            logger.info(f"PDF 解析完成，提取了 {len(full_text)} 个字符")
            return full_text
            
        except Exception as e:
            logger.error(f"解析 PDF 文件失败 {file_path}: {str(e)}")
            return None
    
    @staticmethod
    def parse_docx(file_path: str) -> Optional[str]:
        """
        解析 Word 文档，提取文本内容
        
        Args:
            file_path: Word 文档路径
            
        Returns:
            提取的文本内容，失败时返回 None
        """
        try:
            logger.info(f"开始解析 Word 文档: {file_path}")
            
            doc = DocxDocument(file_path)
            text_content = []
            
            # 提取段落文本
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            full_text = "\n\n".join(text_content)
            logger.info(f"Word 文档解析完成，提取了 {len(full_text)} 个字符")
            return full_text
            
        except Exception as e:
            logger.error(f"解析 Word 文档失败 {file_path}: {str(e)}")
            return None
    
    @classmethod
    def parse(cls, file_path: str, file_type: str) -> Optional[str]:
        """
        根据文件类型解析文档
        
        Args:
            file_path: 文件路径
            file_type: 文件类型 (pdf/docx)
            
        Returns:
            提取的文本内容，失败时返回 None
        """
        file_type = file_type.lower()
        
        if file_type == "pdf":
            return cls.parse_pdf(file_path)
        elif file_type in ["docx", "doc"]:
            return cls.parse_docx(file_path)
        else:
            logger.error(f"不支持的文件类型: {file_type}")
            return None
    
    @staticmethod
    def get_file_type(filename: str) -> str:
        """
        从文件名获取文件类型
        
        Args:
            filename: 文件名
            
        Returns:
            文件类型（不含点号）
        """
        return Path(filename).suffix[1:].lower()
